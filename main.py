import whisper
import cohere
import os
import time
import pygame
from datetime import datetime
from babel.dates import format_date
import warnings
import sounddevice as sd
import queue
import numpy as np
import io
import wave
from gtts import gTTS
import threading  # Importando a biblioteca threading

from tqdm import tqdm
from docx import Document

from tkinter import *
from tkinter import ttk

warnings.filterwarnings(
    "ignore", message="FP16 is not supported on CPU; using FP32 instead")

# Configuração do cliente Cohere
co = cohere.Client('OO9QG57UjoEEwxA0E7x99HELb1fkmM6WIrnXQX64')

# Inicializar o modelo Whisper
modelo = whisper.load_model("base", device="cpu")

# Inicializar o pygame.mixer uma vez
pygame.mixer.init()

# Variável global para armazenar a frase reconhecida
resultado = {"text": ""}

# Variáveis globais e lock para controle de gravação de áudio
audio_lock = threading.Lock()
audio_gravado = None  # Variável global para armazenar o áudio gravado


def gerarTexto(texto):
    """Gera texto formal a partir da entrada usando o Cohere."""
    try:
        response = co.generate(
            model='command-xlarge-nightly',
            prompt=f'Corrija o texto abaixo, ajustando apenas erros ortográficos, gramaticais, de acentuação e de pontuação. \
            Mantenha a estrutura, o estilo e as palavras o mais próximo possível do original, sem que haja substituição por outras palavras parecidas.\
            Criar um cabeçalho, com o tema abordado no texto. Além disso,\
            criar outro cabeçalho, escrevendo o nome da pessoa e seu número de matricula, conforme exemplo: Nome: nome da pessoa - Matrícula: número da matrícula\
            . Pular uma linha e iniciar o texto. Ao final do texto, criar uma linha para assinatura e, abaixo da linha, conter o nome da pessoa e seu número de matrícula.\
            :\n\nTexto: {texto}',
            max_tokens=500,
            temperature=0.7
        )
        return response.generations[0].text.strip()
    except Exception as e:
        print(f"Erro ao gerar texto: {e}")
        return "Erro ao processar o texto."


def gravar_audio(duracao=10, taxa_amostragem=16000):
    """Grava áudio do microfone e retorna como um array de bytes."""
    global audio_gravado
    global fraseFormatada

    print("\nAguarde... Gravando o áudio...\n")

    # Configurar tamanho fixo para a label
    label_resultado.config(text="", width=50, anchor="center")
    label_resultado.update_idletasks()  # Atualiza a interface imediatamente

    # Exibir mensagem de gravação
    label_resultado.config(text="Gravando...", width=50)
    label_resultado.update_idletasks()

    # Usando lock para garantir que a gravação seja realizada sem interferência
    with audio_lock:
        # Gravar o áudio usando o SoundDevice
        audio_gravado = sd.rec(int(duracao * taxa_amostragem),
                               samplerate=taxa_amostragem, channels=1, dtype='int16')
    sd.wait()  # Esperar até a gravação terminar

    # Salvar o áudio gravado em um arquivo temporário
    caminho_audio = "audio_temporario.wav"
    salvar_audio_em_arquivo(audio_gravado, caminho_audio)

    # Exibir mensagem de transcrição
    label_resultado.config(text="Transcrevendo...", width=50)
    label_resultado.update_idletasks()

    frase = transcrever_audio(caminho_audio)
    # EXIBIRÁ A FRASE TRANSCRITA PELO WHISPER
    print(f"\nVocê disse: {frase}\n")

    fraseFormatada = gerarTexto(frase)

    # Exibir o texto transcrito
    resultado["text"] = frase or "Nenhum som detectado."
    label_resultado.config(text=fraseFormatada, width=110)
    label_resultado.update_idletasks()

    return 'sucesso'


def salvar_audio_em_arquivo(audio, caminho_arquivo):
    """Salva o áudio gravado em um arquivo WAV."""
    with wave.open(caminho_arquivo, 'wb') as arquivo_wav:
        arquivo_wav.setnchannels(1)
        arquivo_wav.setsampwidth(2)  # 2 bytes por amostra (16-bit)
        arquivo_wav.setframerate(16000)
        arquivo_wav.writeframes(audio.tobytes())


def transcrever_audio(arquivo_audio):
    """Usa o modelo Whisper para transcrever o áudio."""
    global resultado
    try:
        # Especificando o idioma como português (pt)
        resultado = modelo.transcribe(
            arquivo_audio, language='pt')  # Especificando o idioma

        return resultado["text"]
    except Exception as e:
        print(f"Erro ao transcrever áudio: {e}")
        return None


def reproduzir_resposta(resposta):
    """Gera e reproduz áudio com a resposta do assistente."""
    arquivo_audio = "resposta.mp3"

    # Remover o arquivo se ele existir
    if os.path.exists(arquivo_audio):
        os.remove(arquivo_audio)

    # Criar o arquivo de áudio usando gTTS
    tts = gTTS(text=resposta, lang='pt')
    tts.save(arquivo_audio)

    # Reproduzir o áudio
    pygame.mixer.music.load(arquivo_audio)
    pygame.mixer.music.play()

    # Esperar até que a reprodução seja concluída
    while pygame.mixer.music.get_busy():
        time.sleep(0.1)

    # Liberar o arquivo de áudio
    pygame.mixer.music.unload()

    # Remover o arquivo de áudio
    os.remove(arquivo_audio)


def reconhecer_voz():
    """Captura o áudio com o microfone e usa Whisper para transcrever."""
    global resultado
    try:
        # Configurar tamanho fixo para a label
        label_resultado.config(text="", width=50, anchor="center")
        label_resultado.update_idletasks()  # Atualiza a interface imediatamente

        # Exibir mensagem de gravação
        label_resultado.config(text="Gravando...", width=50)
        label_resultado.update_idletasks()

        audio = gravar_audio()  # 5 segundos de gravação

        # Salvar o áudio gravado em um arquivo temporário
        caminho_audio = "audio_temporario.wav"
        salvar_audio_em_arquivo(audio, caminho_audio)

        # Exibir mensagem de transcrição
        label_resultado.config(text="Transcrevendo...", width=50)
        label_resultado.update_idletasks()

        # Transcrever o áudio com o modelo Whisper
        frase = transcrever_audio(caminho_audio)
        print(f"\nVocê disse: {frase}\n")

        # Exibir o texto transcrito
        resultado["text"] = frase or "Nenhum som detectado."
        label_resultado.config(text=resultado["text"], width=50)
        label_resultado.update_idletasks()

        # Remover o arquivo de áudio temporário após a transcrição
        os.remove(caminho_audio)

        return frase
    except Exception as e:
        print(f"Erro ao reconhecer voz: {e}")
        resultado["text"] = "Erro ao processar áudio."
        label_resultado.config(text=resultado["text"], width=50)
        label_resultado.update_idletasks()
        return None


# Variável global para controle
executando = True


def main():
    """Função principal que executa o loop do assistente."""
    global executando
    global ativacao

    executando = True
    try:
        print("Assistente em execução...")
        ativacao = gravar_audio()
    except KeyboardInterrupt:
        print("\nPrograma interrompido pelo usuário.")
    finally:
        pygame.mixer.quit()


def iniciar_assistente():
    """Inicia o assistente em uma thread separada."""
    thread = threading.Thread(target=main, daemon=True)
    thread.start()


def parar_assistente():
    """Interrompe o loop principal e encerra o programa."""
    global executando
    global ativacao
    executando = False

    print('\nGravação interrompida\n')

    sd.stop()  # INTERROMPE A GRAVAÇÃO


def sair_assistente():  # SAIR E FECHAR O ASSISTENTE
    print("\nSaindo da aplicação...\n")

    root.destroy()


def carregarBarra(m, varBarra, root):
    cont = 0
    etapas = m

    # Exibe a barra de progresso
    progressbar.grid(row=5, column=0, columnspan=4, padx=20, pady=20, sticky='ew')  # Garante que a barra seja visível

    while cont < etapas:
        cont += 1
        time.sleep(1)  # Simula o tempo de progresso
        varBarra.set((cont / etapas) * 100)  # Atualiza a barra em porcentagem
        root.update_idletasks()  # Atualiza a interface sem processar eventos desnecessários

    # Esconde a barra de progresso após a conclusão
    progressbar.grid_forget()


def criarArquivoTexto(texto):
    print("\nAguarde!Criando documento de texto...\n")

    """Função responsável por criar um arquivo Word (.docx) contendo o texto formatado."""
    # Garante que o diretório "arquivos" exista
    os.makedirs("textos", exist_ok=True)

    global fraseFormatada  # VARIÁVEL RECEBERÁ O TEXTO FORMATADO

    try:
        # Cria um novo documento Word
        document = Document()

        texto_formatado = fraseFormatada
        # Captura a data atual no formato desejado
        dataAtual = format_date(
            datetime.now(), format="d 'de' MMMM 'de' yyyy", locale="pt_BR")
        cabecalho = f"Rio de Janeiro, {dataAtual}"
        # Adiciona a data no início do documento
        document.add_paragraph(cabecalho)

        # Divide o texto em linhas
        linhas = texto_formatado.splitlines()
        total_linhas = len(linhas)

        # Adiciona cada linha ao documento com uma barra de progresso do TQDM
        # for linha in tqdm(linhas, total=total_linhas, desc="Criando arquivo Word", ascii=True, ncols=80, unit="linha"):
        #     document.add_paragraph(linha)  # Adiciona uma linha como parágrafo
        #     time.sleep(0.5)  # Simula o tempo de escrita

        # Inicializa a barra de progresso
        carregarBarra(total_linhas, varBarra, root)
        document.add_paragraph(linhas)  # Adiciona uma linha como parágrafo
        time.sleep(0.5)  # Simula o tempo de escrita

        # Salva o arquivo Word
        caminho_arquivo = "textos/texto.docx"
        document.save(caminho_arquivo)

        print(f"\nArquivo Word criado com sucesso em: {caminho_arquivo}")
    except Exception as e:
        print(f"Erro ao criar o arquivo Word: {e}")


# Função para centralizar a janela
def centralizar_janela(janela, largura, altura):
    screen_width = janela.winfo_screenwidth()
    screen_height = janela.winfo_screenheight()
    pos_x = (screen_width // 2) - (largura // 2)
    pos_y = (screen_height // 2) - (altura // 2)
    janela.geometry(f"{largura}x{altura}+{pos_x}+{pos_y}")


# Configuração da interface gráfica com Tkinter
root = Tk()
largura_janela = 800
altura_janela = 800
centralizar_janela(root, largura_janela, altura_janela)

frm = ttk.Frame(root, width=largura_janela, height=altura_janela)
frm.grid()
frm.grid_propagate(False)  # Impedir que o Frame redimensione com base nos widgets
# Elementos da interface
ttk.Label(frm, text="Assistente de Voz").grid(column=0,columnspan=6, row=0, sticky='w')

# Label para exibir o texto reconhecido
label_resultado = ttk.Label(frm, text=f'{resultado["text"]}', wraplength=frm.winfo_width() - 40)  # Ajustar wraplength com base na largura do Frame
label_resultado.grid(column=0, row=4, sticky='ew', padx=20)

varBarra = DoubleVar()

progressbar = ttk.Progressbar(frm, orient='horizontal', mode="determinate", variable=varBarra,
                              style='success.Striped.Horizontal.TProgressbar', maximum=100)
progressbar.grid(column=0, row=5, columnspan=4, padx=20, pady=20, sticky='ew')
progressbar.grid_forget()  # Inicialmente oculta a barra de progresso


# Botões
ttk.Button(frm, text="Iniciar Assistente", command=iniciar_assistente).grid(column=0, row=1, sticky='ew')
ttk.Button(frm, text="Parar Assistente", command=parar_assistente).grid(column=1, row=1, sticky='ew')
ttk.Button(frm, text="Sair", command=sair_assistente).grid(column=2, row=1, sticky='ew')
ttk.Button(frm, text="Gerar Arquivo", command=lambda: criarArquivoTexto(resultado['text'])).grid(column=0, row=3, sticky='ew')

root.mainloop()
